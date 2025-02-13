import os
from openai import OpenAI
import docx
import time

# Cấu hình API Key
client = OpenAI(
    base_url="https://api.groq.com/openai/v1",
    api_key='xxx',
)
# Lấy nội dung từ file docx
def read_docx_file(file_path):
    doc = docx.Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()!=""])
    return full_text

# Chia nhỏ nội dung
def split_text(text, max_chars):
    chunks = []
    while len(text) > max_chars:
        split_pos = text.rfind("\n", 0, max_chars)
        if split_pos == -1:
            split_pos = max_chars
        chunk = text[:split_pos]
        chunks.append(chunk.strip())
        text = text[split_pos:].strip()
    if text:
        chunks.append(text.strip())
    return chunks

# Dịch nội dung
def translate_text(text, source_lang="Tiếng Anh", target_lang="Tiếng Việt"):
    prompt = f"Hãy dịch nội dung sau từ {source_lang} sang {target_lang}:<{text}>"
    messages = [
        {
            "role": "system",
            "content": "Bạn là một thông dịch viên chuyên nghiệp."
        },
        {
            "role": "user",
            "content": prompt
        }
    ]
    response = client.chat.completions.create(
        messages=messages,
        model="gemma2-9b-it",
        temperature=0.3,
        max_tokens=1024
    )
    translation = response.choices[0].message.content.strip()
    return translation

# Tạo file docx lưu kết quả dịch thuật
def write_docx_file(text, file_path):
    doc = docx.Document()
    for t in text:
        doc.add_paragraph(t)
    doc.save(file_path)

def main():
    input_file = "input.docx"
    output_file = "output.docx"
    input_text = read_docx_file(input_file)
    if not input_text:
        print("KHÔNG TÌM THẤY NỘI DUNG FILE")
        return
    max_tokens = 500
    max_chars = max_tokens*4
    text_chunks = split_text(input_text, max_chars)

    translate_chunks = []
    for idx, chunk in enumerate(text_chunks):
        print(f"Đang dịch phần {idx+1}/{len(text_chunks)} ...")
        translated = translate_text(chunk)
        if translated:
            translate_chunks.append(translated)
        else:
            print(f"Lỗi khi dịch phần {idx+1}.")
    final_translation = "\n\n".join(translate_chunks)
    paragraphs = final_translation.split("\n")
    write_docx_file(paragraphs, output_file)
    print(f"Quá trình dịch hoàn tất. File dịch được lưu tại: {output_file}")

if __name__ == "__main__":
    main()
